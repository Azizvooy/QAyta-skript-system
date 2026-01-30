from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from datetime import datetime
from pathlib import Path
import sys

# Базовые пути
BASE_DIR = Path(__file__).parent.parent.parent
DATA_DIR = BASE_DIR / 'data'
OUTPUT_DIR = BASE_DIR / 'output' / 'reports'

# Создаем директории
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Создаем документ
doc = Document()

# Заголовок
title = doc.add_heading('ОТЧЕТ ПО РЕЗУЛЬТАТАМ ОБЗВОНА ЗАЯВОК', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Дата отчета
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run(f'Дата формирования отчета: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
date_run.font.size = Pt(11)

doc.add_paragraph()

# Загружаем данные
print("Загрузка данных...")
file_path = DATA_DIR / 'ALL_DATA_CLEANED.csv'

if not file_path.exists():
    print(f"❌ ОШИБКА: Файл не найден: {file_path}")
    sys.exit(1)

try:
    df = pd.read_csv(file_path, encoding='utf-8-sig', low_memory=False)
except Exception as e:
    print(f"❌ ОШИБКА при загрузке: {e}")
    sys.exit(1)

# Конвертируем даты
df['Дата фиксации'] = pd.to_datetime(df['Дата фиксации'], format='%d.%m.%Y %H:%M:%S', errors='coerce')
df_2025 = df[df['Дата фиксации'].dt.year == 2025].copy()

# ========== ОБЩАЯ СТАТИСТИКА ==========
doc.add_heading('1. ОБЩАЯ СТАТИСТИКА', 1)

p = doc.add_paragraph()
p.add_run(f'📊 Всего записей звонков за 2025 год: ').bold = True
p.add_run(f'{len(df_2025):,}')

p = doc.add_paragraph()
p.add_run(f'📋 Уникальных заявок: ').bold = True
p.add_run(f'{df_2025["Номер карты"].nunique():,}')

doc.add_paragraph()

# ========== РАСПРЕДЕЛЕНИЕ ПО МЕСЯЦАМ ==========
doc.add_heading('2. РАСПРЕДЕЛЕНИЕ ПО МЕСЯЦАМ', 1)

df_2025['Месяц_число'] = df_2025['Дата фиксации'].dt.month
months_dict = {9: 'Сентябрь', 10: 'Октябрь', 11: 'Ноябрь', 12: 'Декабрь'}

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Месяц'
hdr_cells[1].text = 'Количество звонков'
hdr_cells[2].text = 'Уникальных заявок'

for month_num in sorted(df_2025['Месяц_число'].unique()):
    month_data = df_2025[df_2025['Месяц_число'] == month_num]
    row_cells = table.add_row().cells
    row_cells[0].text = months_dict.get(month_num, str(month_num))
    row_cells[1].text = f'{len(month_data):,}'
    row_cells[2].text = f'{month_data["Номер карты"].nunique():,}'

doc.add_paragraph()

# ========== ДЕТАЛЬНАЯ СТАТИСТИКА РЕЗУЛЬТАТОВ ==========
doc.add_heading('3. РЕЗУЛЬТАТЫ ЗВОНКОВ', 1)

positive = df_2025[df_2025['Статус'].str.lower().str.contains('положит', na=False)]
negative = df_2025[df_2025['Статус'].str.lower().str.contains('отрицат', na=False)]
no_answer = df_2025[df_2025['Статус'].str.lower().str.contains('нет ответа|занято', na=False)]
closed = df_2025[df_2025['Статус'].str.lower().str.contains('закрыта', na=False)]
disconnected = df_2025[df_2025['Статус'].str.lower().str.contains('соед.прервано|прервано', na=False)]
silence = df_2025[df_2025['Статус'].str.lower().str.contains('тишина', na=False)]
medical = df_2025[df_2025['Статус'].str.lower().str.contains('тиббиёт|ходими', na=False)]

# Таблица результатов
table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Категория'
hdr_cells[1].text = 'Количество'
hdr_cells[2].text = 'Процент'

def add_result_row(table, category, count, total):
    row = table.add_row().cells
    row[0].text = category
    row[1].text = f'{count:,}'
    row[2].text = f'{count/total*100:.2f}%'

add_result_row(table2, '✅ Положительные ответы', len(positive), len(df_2025))
add_result_row(table2, '❌ Отрицательные ответы', len(negative), len(df_2025))
add_result_row(table2, '📞 Нет ответа/Занято', len(no_answer), len(df_2025))
add_result_row(table2, '🚫 Заявка закрыта', len(closed), len(df_2025))
add_result_row(table2, '📵 Соединение прервано', len(disconnected), len(df_2025))
add_result_row(table2, '🔇 Тишина', len(silence), len(df_2025))
add_result_row(table2, '🏥 Медработники', len(medical), len(df_2025))

doc.add_paragraph()

# ========== ИТОГОВАЯ СВОДКА ==========
doc.add_heading('4. ИТОГОВАЯ СВОДКА', 1)

total_reached = len(positive) + len(negative)
total_not_reached = len(no_answer) + len(closed) + len(disconnected) + len(silence)

p = doc.add_paragraph()
p.add_run('Дозвонились и получили ответ: ').bold = True
p.add_run(f'{total_reached:,} ({total_reached/len(df_2025)*100:.2f}%)')

p = doc.add_paragraph(style='List Bullet')
if total_reached > 0:
    p.add_run(f'✅ Положительных: {len(positive):,} ({len(positive)/total_reached*100:.2f}% от дозвонившихся)')
else:
    p.add_run(f'✅ Положительных: {len(positive):,} (0.00% от дозвонившихся)')

p = doc.add_paragraph(style='List Bullet')
if total_reached > 0:
    p.add_run(f'❌ Отрицательных: {len(negative):,} ({len(negative)/total_reached*100:.2f}% от дозвонившихся)')
else:
    p.add_run(f'❌ Отрицательных: {len(negative):,} (0.00% от дозвонившихся)')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('НЕ дозвонились: ').bold = True
p.add_run(f'{total_not_reached:,} ({total_not_reached/len(df_2025)*100:.2f}%)')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'📞 Нет ответа/Занято: {len(no_answer):,}')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'🚫 Заявка закрыта: {len(closed):,}')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'📵 Соединение прервано: {len(disconnected):,}')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'🔇 Тишина: {len(silence):,}')

doc.add_paragraph()

# ========== ТОП ОПЕРАТОРОВ ==========
doc.add_heading('5. ТОП-10 ОПЕРАТОРОВ ПО КОЛИЧЕСТВУ ЗВОНКОВ', 1)

top_operators = df_2025['Оператор'].value_counts().head(10)

table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Light Grid Accent 1'
hdr_cells = table3.rows[0].cells
hdr_cells[0].text = '№'
hdr_cells[1].text = 'Оператор'
hdr_cells[2].text = 'Количество звонков'

for i, (operator, count) in enumerate(top_operators.items(), 1):
    row = table3.add_row().cells
    row[0].text = str(i)
    row[1].text = operator
    row[2].text = f'{count:,}'

doc.add_paragraph()

# ========== ЭФФЕКТИВНОСТЬ ==========
doc.add_heading('6. ПОКАЗАТЕЛИ ЭФФЕКТИВНОСТИ', 1)

efficiency = len(positive) / total_reached * 100 if total_reached > 0 else 0
conversion = len(positive) / len(df_2025) * 100

p = doc.add_paragraph()
p.add_run('Конверсия в положительный результат:').bold = True

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'От всех звонков: {conversion:.2f}%')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'От дозвонившихся: {efficiency:.2f}%')

# Сохраняем документ
timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
output_file = OUTPUT_DIR / f'Отчет_по_обзвону_2025_{timestamp}.docx'

try:
    doc.save(output_file)
    print(f"\n✅ Документ успешно создан: {output_file}")
    print(f"📄 Отчет содержит полную статистику по обзвону за 2025 год")
except Exception as e:
    print(f"\n❌ ОШИБКА при сохранении: {e}")
    sys.exit(1)
