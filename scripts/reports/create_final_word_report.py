from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path
import sys

# Базовые пути
BASE_DIR = Path(__file__).parent.parent.parent
OUTPUT_DIR = BASE_DIR / 'output' / 'reports'

# Создаем директории
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Создаем документ
doc = Document()

# Заголовок
title = doc.add_heading('ПОЛНЫЙ ОТЧЕТ ПО ОБЗВОНУ ЗАЯВОК', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Дата отчета
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
date_run.font.size = Pt(11)

doc.add_paragraph()

# Данные из анализа
total = 753850
positive = 345720
negative = 47609
no_answer = 189657
closed = 159199
disconnected = 1195
silence = 111
medical = 9972
other = 108
unknown = 279

dozonil = positive + negative
ne_dozonil = no_answer + closed + disconnected + silence

# ========== ОБЩАЯ СТАТИСТИКА ==========
doc.add_heading('1. ОБЩАЯ СТАТИСТИКА', 1)

p = doc.add_paragraph()
p.add_run(f'📊 Всего записей звонков: ').bold = True
p.add_run(f'1,482,281')

p = doc.add_paragraph()
p.add_run(f'🎫 Уникальных заявок (карт): ').bold = True
run = p.add_run(f'{total:,}')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 102, 204)
run.bold = True

doc.add_paragraph()

# ========== РЕЗУЛЬТАТЫ ОБЗВОНА ==========
doc.add_heading('2. РЕЗУЛЬТАТЫ ОБЗВОНА ПО УНИКАЛЬНЫМ ЗАЯВКАМ', 1)

# Таблица результатов
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Категория'
hdr_cells[1].text = 'Количество'
hdr_cells[2].text = 'Процент'

def add_row(table, category, count, total):
    row = table.add_row().cells
    row[0].text = category
    row[1].text = f'{count:,}'
    row[2].text = f'{count/total*100:.2f}%'

add_row(table, '✅ Положительные', positive, total)
add_row(table, '❌ Отрицательные', negative, total)
add_row(table, '📞 Нет ответа/Занято', no_answer, total)
add_row(table, '🚫 Заявка закрыта', closed, total)
add_row(table, '📵 Соединение прервано', disconnected, total)
add_row(table, '🔇 Тишина', silence, total)
add_row(table, '🏥 Медработники', medical, total)
add_row(table, '📝 Прочее', other, total)
add_row(table, '❓ Неизвестно', unknown, total)

doc.add_paragraph()

# ========== ИТОГОВАЯ СВОДКА ==========
doc.add_heading('3. ИТОГОВАЯ СВОДКА', 1)

p = doc.add_paragraph()
p.add_run('Дозвонились и получили ответ: ').bold = True
run = p.add_run(f'{dozonil:,} ({dozonil/total*100:.2f}%)')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 176, 80)
run.bold = True

p = doc.add_paragraph(style='List Bullet')
run = p.add_run(f'✅ Положительных: {positive:,} ({positive/dozonil*100:.2f}% от дозвонившихся)')
run.font.size = Pt(11)

p = doc.add_paragraph(style='List Bullet')
run = p.add_run(f'❌ Отрицательных: {negative:,} ({negative/dozonil*100:.2f}% от дозвонившихся)')
run.font.size = Pt(11)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('НЕ дозвонились: ').bold = True
run = p.add_run(f'{ne_dozonil:,} ({ne_dozonil/total*100:.2f}%)')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(192, 0, 0)
run.bold = True

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'📞 Нет ответа/Занято: {no_answer:,}')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'🚫 Заявка закрыта: {closed:,}')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'📵 Соединение прервано: {disconnected:,}')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'🔇 Тишина: {silence:,}')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Медработники: ').bold = True
p.add_run(f'{medical:,} ({medical/total*100:.2f}%)')

doc.add_paragraph()

# ========== ПОКАЗАТЕЛИ ЭФФЕКТИВНОСТИ ==========
doc.add_heading('4. ПОКАЗАТЕЛИ ЭФФЕКТИВНОСТИ', 1)

efficiency = positive / dozonil * 100
conversion = positive / total * 100

p = doc.add_paragraph()
p.add_run('Конверсия в положительный результат:').bold = True

p = doc.add_paragraph(style='List Bullet')
run = p.add_run(f'От всех заявок: {conversion:.2f}%')
run.font.size = Pt(12)
run.bold = True

p = doc.add_paragraph(style='List Bullet')
run = p.add_run(f'От дозвонившихся: {efficiency:.2f}%')
run.font.size = Pt(12)
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

p = doc.add_paragraph()
p.add_run('Общая эффективность:').bold = True

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'Успешных дозвонов: {dozonil/total*100:.2f}%')

# ========== ТОП ОПЕРАТОРОВ ==========
doc.add_heading('5. ТОП-20 ОПЕРАТОРОВ', 1)

operators = [
    ('Narziyeva Gavxar Atxamjanovna', 52641),
    ('Xusniddinova Shaxnoza Akramovna', 51137),
    ('Sagdullayeva Moxinur Asqar qizi', 47857),
    ('Muxamadaliyeva Mufazzal Abduqaxxorovna', 44643),
    ('Rahimjonov Kamoliddin Olimjon o\'g\'li', 44547),
    ('Mirbabayeva Shirin Kaxramonovna', 37939),
    ('Karimova Durdona Toir qizi', 36599),
    ('Payziyeva Shoxista Navro\'zjon qizi', 35792),
    ('Zokirjonova Surayyo Rustam qizi', 35480),
    ('Umarova Madina Talat qizi', 35414),
    ('Xoshimov Akromjon Axmadjon o\'g\'li', 29662),
    ('Sobirjonova Umidaxon Rustamovna', 29352),
    ('Xasanova Maftuna Askar qizi', 24803),
    ('Niyozova Nigora Abduvali qizi', 21957),
    ('Qosimov Firdavs Nuriddin o\'g\'li', 18249),
    ('Mavlyanova Dilobar Rustam qizi', 18070),
    ('Sirojiddinov Ismoilbek Shavkat o\'g\'li', 17380),
    ('Abduraximov Otabek Abdumanapovich', 16568),
    ('Mahmudjonova Mushtari Akram qizi', 9183),
    ('Zokirjonova Saida Murodjon qizi', 7837),
]

table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Light Grid Accent 1'
hdr_cells = table3.rows[0].cells
hdr_cells[0].text = '№'
hdr_cells[1].text = 'Оператор'
hdr_cells[2].text = 'Количество заявок'

for idx, (op, count) in enumerate(operators, 1):
    row = table3.add_row().cells
    row[0].text = str(idx)
    row[1].text = op
    row[2].text = f'{count:,}'

p = doc.add_paragraph()
p.add_run('\nВсего уникальных операторов: 48').italic = True

doc.add_paragraph()

# ========== КЛЮЧЕВЫЕ ПОКАЗАТЕЛИ ==========
doc.add_heading('6. КЛЮЧЕВЫЕ ПОКАЗАТЕЛИ', 1)

key_stats = [
    ('Общее количество уникальных заявок', f'{total:,}'),
    ('Успешных положительных результатов', f'{positive:,} (45.86%)'),
    ('Процент конверсии от дозвонившихся', f'{efficiency:.2f}%'),
    ('Количество дозвонов', f'{dozonil:,} (52.18%)'),
    ('Эффективность работы операторов', 'Высокая')
]

table4 = doc.add_table(rows=1, cols=2)
table4.style = 'Light Grid Accent 1'
hdr_cells = table4.rows[0].cells
hdr_cells[0].text = 'Показатель'
hdr_cells[1].text = 'Значение'

for stat_name, stat_value in key_stats:
    row = table4.add_row().cells
    row[0].text = stat_name
    row[1].text = stat_value

# Сохраняем документ
timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
output_file = OUTPUT_DIR / f'ФИНАЛЬНЫЙ_ОТЧЕТ_{timestamp}.docx'

try:
    doc.save(output_file)
    print(f"\n✅ ОТЧЕТ СОЗДАН: {output_file}")
    print(f"\n📊 ИТОГОВАЯ СТАТИСТИКА:")
    print(f"   🎫 Уникальных заявок: {total:,}")
    print(f"   ✅ Положительных: {positive:,} ({positive/total*100:.2f}%)")
    print(f"   ❌ Отрицательных: {negative:,} ({negative/total*100:.2f}%)")
    print(f"   🚫 Закрытых: {closed:,} ({closed/total*100:.2f}%)")
    print(f"   📞 Дозвонились: {dozonil:,} ({dozonil/total*100:.2f}%)")
    if dozonil > 0:
        print(f"   🎯 Конверсия: {efficiency:.2f}% (от дозвонившихся)")
except Exception as e:
    print(f"\n❌ ОШИБКА при сохранении: {e}")
    sys.exit(1)

print(f"✅ ОТЧЕТ СОЗДАН: {output_file}")
print(f"\n📊 ИТОГОВАЯ СТАТИСТИКА:")
print(f"   🎫 Уникальных заявок: {total:,}")
print(f"   ✅ Положительных: {positive:,} ({positive/total*100:.2f}%)")
print(f"   ❌ Отрицательных: {negative:,} ({negative/total*100:.2f}%)")
print(f"   🚫 Закрытых: {closed:,} ({closed/total*100:.2f}%)")
print(f"   📞 Дозвонились: {dozonil:,} ({dozonil/total*100:.2f}%)")
print(f"   🎯 Конверсия: {efficiency:.2f}% (от дозвонившихся)")
