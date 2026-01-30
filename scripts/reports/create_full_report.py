from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from datetime import datetime
from pathlib import Path
import sys

# Базовые пути
BASE_DIR = Path(__file__).parent.parent.parent
DATA_DIR = BASE_DIR / 'data'
OUTPUT_DIR = BASE_DIR / 'output' / 'reports'

# Создаем директории
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Создаем документ
doc = Document()

# Заголовок
title = doc.add_heading('ПОЛНЫЙ ОТЧЕТ ПО ОБЗВОНУ ЗАЯВОК', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Дата отчета
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
date_run.font.size = Pt(11)

doc.add_paragraph()

# Загружаем данные
print("Загрузка данных...")
file_path = DATA_DIR / 'ALL_DATA_COLLECTED.csv'

if not file_path.exists():
    print(f"❌ ОШИБКА: Файл не найден: {file_path}")
    sys.exit(1)

try:
    df = pd.read_csv(file_path, encoding='utf-8-sig', low_memory=False)
    print(f"Всего записей: {len(df):,}")
except Exception as e:
    print(f"❌ ОШИБКА при загрузке: {e}")
    sys.exit(1)

# Определяем колонки
номер_карты = 'Номер карты'
статус = 'Статус связи'
оператор = 'Оператор'

# Фильтруем записи с номером карты
df_clean = df[df[номер_карты].notna()].copy()
print(f"Записей с номером карты: {len(df_clean):,}")

# Берем последний статус для каждой уникальной карты
df_unique = df_clean.drop_duplicates(subset=номер_карты, keep='last')
print(f"Уникальных карт: {len(df_unique):,}")

# ========== ОБЩАЯ СТАТИСТИКА ==========
doc.add_heading('1. ОБЩАЯ СТАТИСТИКА', 1)

p = doc.add_paragraph()
p.add_run(f'📊 Всего записей звонков: ').bold = True
p.add_run(f'{len(df_clean):,}')

p = doc.add_paragraph()
p.add_run(f'🎫 Уникальных заявок (карт): ').bold = True
p.add_run(f'{len(df_unique):,}')

doc.add_paragraph()

# ========== РЕЗУЛЬТАТЫ ОБЗВОНА ==========
doc.add_heading('2. РЕЗУЛЬТАТЫ ОБЗВОНА', 1)

# Классификация
positive = df_unique[df_unique[статус].astype(str).str.lower().str.contains('положит', na=False)]
negative = df_unique[df_unique[статус].astype(str).str.lower().str.contains('отрицат', na=False)]
no_answer = df_unique[df_unique[статус].astype(str).str.lower().str.contains('нет ответа|занято', na=False)]
closed = df_unique[df_unique[статус].astype(str).str.lower().str.contains('закрыта', na=False)]
disconnected = df_unique[df_unique[статус].astype(str).str.lower().str.contains('соед.прервано|прервано', na=False)]
silence = df_unique[df_unique[статус].astype(str).str.lower().str.contains('тишина', na=False)]
medical = df_unique[df_unique[статус].astype(str).str.lower().str.contains('тиббиёт|ходими', na=False)]

total = len(df_unique)

# Таблица результатов
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Категория'
hdr_cells[1].text = 'Количество'
hdr_cells[2].text = 'Процент'

def add_row(table, category, count, total):
    row = table.add_row().cells
    row[0].text = category
    row[1].text = f'{count:,}'
    row[2].text = f'{count/total*100:.2f}%'

add_row(table, '✅ Положительные', len(positive), total)
add_row(table, '❌ Отрицательные', len(negative), total)
add_row(table, '📞 Нет ответа/Занято', len(no_answer), total)
add_row(table, '🚫 Заявка закрыта', len(closed), total)
add_row(table, '📵 Соединение прервано', len(disconnected), total)
add_row(table, '🔇 Тишина', len(silence), total)
add_row(table, '🏥 Медработники', len(medical), total)

doc.add_paragraph()

# ========== ИТОГОВАЯ СВОДКА ==========
doc.add_heading('3. ИТОГОВАЯ СВОДКА', 1)

dozonil = len(positive) + len(negative)
ne_dozonil = len(no_answer) + len(closed) + len(disconnected) + len(silence)

p = doc.add_paragraph()
p.add_run('Дозвонились и получили ответ: ').bold = True
p.add_run(f'{dozonil:,} ({dozonil/total*100:.2f}%)')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'✅ Положительных: {len(positive):,} ({len(positive)/dozonil*100:.2f}% от дозвонившихся)')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'❌ Отрицательных: {len(negative):,} ({len(negative)/dozonil*100:.2f}% от дозвонившихся)')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('НЕ дозвонились: ').bold = True
p.add_run(f'{ne_dozonil:,} ({ne_dozonil/total*100:.2f}%)')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'📞 Нет ответа/Занято: {len(no_answer):,}')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'🚫 Заявка закрыта: {len(closed):,}')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'📵 Соединение прервано: {len(disconnected):,}')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'🔇 Тишина: {len(silence):,}')

doc.add_paragraph()

# ========== ДЕТАЛЬНАЯ СТАТИСТИКА СТАТУСОВ ==========
doc.add_heading('4. ДЕТАЛЬНАЯ СТАТИСТИКА СТАТУСОВ', 1)

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = '№'
hdr_cells[1].text = 'Статус'
hdr_cells[2].text = 'Количество'

for idx, (st, count) in enumerate(df_unique[статус].value_counts().head(15).items(), 1):
    row = table2.add_row().cells
    row[0].text = str(idx)
    row[1].text = str(st)
    row[2].text = f'{count:,} ({count/total*100:.2f}%)'

doc.add_paragraph()

# ========== ТОП ОПЕРАТОРОВ ==========
doc.add_heading('5. ТОП-15 ОПЕРАТОРОВ', 1)

# Убираем пустые и "-"
df_ops = df_unique[df_unique[оператор].notna() & (df_unique[оператор] != '-')]

if len(df_ops) > 0:
    table3 = doc.add_table(rows=1, cols=3)
    table3.style = 'Light Grid Accent 1'
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Оператор'
    hdr_cells[2].text = 'Количество заявок'

    for idx, (op, count) in enumerate(df_ops[оператор].value_counts().head(15).items(), 1):
        row = table3.add_row().cells
        row[0].text = str(idx)
        row[1].text = op
        row[2].text = f'{count:,}'

doc.add_paragraph()

# ========== ЭФФЕКТИВНОСТЬ ==========
doc.add_heading('6. ПОКАЗАТЕЛИ ЭФФЕКТИВНОСТИ', 1)

efficiency = len(positive) / dozonil * 100 if dozonil > 0 else 0
conversion = len(positive) / total * 100

p = doc.add_paragraph()
p.add_run('Конверсия в положительный результат:').bold = True

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'От всех заявок: {conversion:.2f}%')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'От дозвонившихся: {efficiency:.2f}%')

p = doc.add_paragraph()
p.add_run('Общая эффективность:').bold = True

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'Успешных дозвонов: {dozonil/total*100:.2f}%')

# Сохраняем документ
timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
output_file = OUTPUT_DIR / f'Полный_отчет_по_обзвону_{timestamp}.docx'

try:
    doc.save(output_file)
    print(f"\n✅ Отчет создан: {output_file}")
    print(f"\n📊 КРАТКАЯ СТАТИСТИКА:")
    print(f"   Уникальных заявок: {len(df_unique):,}")
    print(f"   ✅ Положительных: {len(positive):,} ({len(positive)/total*100:.2f}%)")
    print(f"   ❌ Отрицательных: {len(negative):,} ({len(negative)/total*100:.2f}%)")
    print(f"   🚫 Закрытых: {len(closed):,} ({len(closed)/total*100:.2f}%)")
except Exception as e:
    print(f"\n❌ ОШИБКА при сохранении: {e}")
    sys.exit(1)
