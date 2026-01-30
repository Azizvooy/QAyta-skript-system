from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from datetime import datetime

# Создаем документ
doc = Document()

# Заголовок
title = doc.add_heading('ПОЛНЫЙ ОТЧЕТ ПО ОБЗВОНУ ЗАЯВОК', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Дата отчета
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
date_run.font.size = Pt(11)

doc.add_paragraph()

# Загружаем данные
print("📂 Загрузка данных...")
file_path = r'c:\Users\a.djurayev\Desktop\QAyta skript\ALL_DATA_COLLECTED.csv'
df = pd.read_csv(file_path, encoding='utf-8-sig', low_memory=False)

# Определяем колонки
номер_карты = 'Номер карты'
статус = 'Статус связи'
оператор = 'Оператор'

# Фильтруем записи с номером карты
df_clean = df[df[номер_карты].notna()].copy()
print(f"📋 Записей с номером карты: {len(df_clean):,}")

# Берем последний статус для каждой уникальной карты
df_unique = df_clean.drop_duplicates(subset=номер_карты, keep='last')
print(f"🎫 Уникальных заявок: {len(df_unique):,}")

# Функция классификации
def classify_status(status_text):
    if pd.isna(status_text):
        return 'Неизвестно'
    
    text = str(status_text).lower().strip()
    
    if 'положит' in text:
        return 'Положительный'
    elif 'отрицат' in text:
        return 'Отрицательный'
    elif 'нет ответа' in text or 'занято' in text:
        return 'Нет ответа/Занято'
    elif 'закрыта' in text or 'закрыт' in text:
        return 'Заявка закрыта'
    elif 'соед' in text or 'прервано' in text:
        return 'Соединение прервано'
    elif 'тишина' in text:
        return 'Тишина'
    elif 'тиббиёт' in text or 'ходими' in text:
        return 'Медработники'
    else:
        return 'Прочее'

df_unique['Категория'] = df_unique[статус].apply(classify_status)
categories = df_unique['Категория'].value_counts()

total = len(df_unique)

# ========== ОБЩАЯ СТАТИСТИКА ==========
doc.add_heading('1. ОБЩАЯ СТАТИСТИКА', 1)

p = doc.add_paragraph()
p.add_run(f'📊 Всего записей звонков: ').bold = True
p.add_run(f'{len(df_clean):,}')

p = doc.add_paragraph()
p.add_run(f'🎫 Уникальных заявок (карт): ').bold = True
run = p.add_run(f'{total:,}')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 102, 204)

doc.add_paragraph()

# ========== РЕЗУЛЬТАТЫ ОБЗВОНА ==========
doc.add_heading('2. РЕЗУЛЬТАТЫ ОБЗВОНА ПО УНИКАЛЬНЫМ ЗАЯВКАМ', 1)

positive = categories.get('Положительный', 0)
negative = categories.get('Отрицательный', 0)
no_answer = categories.get('Нет ответа/Занято', 0)
closed = categories.get('Заявка закрыта', 0)
disconnected = categories.get('Соединение прервано', 0)
silence = categories.get('Тишина', 0)
medical = categories.get('Медработники', 0)
other = categories.get('Прочее', 0)
unknown = categories.get('Неизвестно', 0)

# Таблица результатов
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Категория'
hdr_cells[1].text = 'Количество'
hdr_cells[2].text = 'Процент'

def add_row(table, category, count, total):
    row = table.add_row().cells
    row[0].text = category
    row[1].text = f'{count:,}'
    row[2].text = f'{count/total*100:.2f}%'

add_row(table, '✅ Положительные', positive, total)
add_row(table, '❌ Отрицательные', negative, total)
add_row(table, '📞 Нет ответа/Занято', no_answer, total)
add_row(table, '🚫 Заявка закрыта', closed, total)
add_row(table, '📵 Соединение прервано', disconnected, total)
add_row(table, '🔇 Тишина', silence, total)
add_row(table, '🏥 Медработники', medical, total)
if other > 0:
    add_row(table, '📝 Прочее', other, total)
if unknown > 0:
    add_row(table, '❓ Неизвестно', unknown, total)

doc.add_paragraph()

# ========== ИТОГОВАЯ СВОДКА ==========
doc.add_heading('3. ИТОГОВАЯ СВОДКА', 1)

dozonil = positive + negative
ne_dozonil = no_answer + closed + disconnected + silence

p = doc.add_paragraph()
p.add_run('Дозвонились и получили ответ: ').bold = True
run = p.add_run(f'{dozonil:,} ({dozonil/total*100:.2f}%)')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 176, 80)

p = doc.add_paragraph(style='List Bullet')
if dozonil > 0:
    p.add_run(f'✅ Положительных: {positive:,} ({positive/dozonil*100:.2f}% от дозвонившихся)')
else:
    p.add_run(f'✅ Положительных: {positive:,} (0.00% от дозвонившихся)')

p = doc.add_paragraph(style='List Bullet')
if dozonil > 0:
    p.add_run(f'❌ Отрицательных: {negative:,} ({negative/dozonil*100:.2f}% от дозвонившихся)')
else:
    p.add_run(f'❌ Отрицательных: {negative:,} (0.00% от дозвонившихся)')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('НЕ дозвонились: ').bold = True
run = p.add_run(f'{ne_dozonil:,} ({ne_dozonil/total*100:.2f}%)')
run.font.color.rgb = RGBColor(192, 0, 0)

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'📞 Нет ответа/Занято: {no_answer:,}')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'🚫 Заявка закрыта: {closed:,}')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'📵 Соединение прервано: {disconnected:,}')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'🔇 Тишина: {silence:,}')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Медработники: ').bold = True
p.add_run(f'{medical:,} ({medical/total*100:.2f}%)')

doc.add_paragraph()

# ========== ПОКАЗАТЕЛИ ЭФФЕКТИВНОСТИ ==========
doc.add_heading('4. ПОКАЗАТЕЛИ ЭФФЕКТИВНОСТИ', 1)

efficiency = positive / dozonil * 100 if dozonil > 0 else 0
conversion = positive / total * 100

p = doc.add_paragraph()
p.add_run('Конверсия в положительный результат:').bold = True

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'От всех заявок: {conversion:.2f}%')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'От дозвонившихся: {efficiency:.2f}%')

p = doc.add_paragraph()
p.add_run('Общая эффективность:').bold = True

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'Успешных дозвонов: {dozonil/total*100:.2f}%')

# ========== ТОП ОПЕРАТОРОВ ==========
doc.add_heading('5. ТОП-15 ОПЕРАТОРОВ', 1)

# Убираем пустые и "-"
df_ops = df_unique[df_unique[оператор].notna() & (df_unique[оператор] != '-')]

if len(df_ops) > 0:
    table3 = doc.add_table(rows=1, cols=3)
    table3.style = 'Light Grid Accent 1'
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Оператор'
    hdr_cells[2].text = 'Количество заявок'

    for idx, (op, count) in enumerate(df_ops[оператор].value_counts().head(15).items(), 1):
        row = table3.add_row().cells
        row[0].text = str(idx)
        row[1].text = op
        row[2].text = f'{count:,}'

doc.add_paragraph()

# ========== ДЕТАЛЬНАЯ СТАТИСТИКА СТАТУСОВ ==========
doc.add_heading('6. ДЕТАЛЬНАЯ СТАТИСТИКА СТАТУСОВ', 1)

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = '№'
hdr_cells[1].text = 'Статус'
hdr_cells[2].text = 'Количество'

for idx, (st, count) in enumerate(df_unique[статус].value_counts().head(15).items(), 1):
    row = table2.add_row().cells
    row[0].text = str(idx)
    row[1].text = str(st)
    row[2].text = f'{count:,} ({count/total*100:.2f}%)'

# Сохраняем документ
timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
output_file = OUTPUT_DIR / f'ИТОГОВЫЙ_ОТЧЕТ_ПО_ОБЗВОНУ_{timestamp}.docx'

try:
    doc.save(output_file)
    print(f"\n✅ ОТЧЕТ СОЗДАН: {output_file}")
    print(f"\n📊 ИТОГОВАЯ СТАТИСТИКА:")
    print(f"   🎫 Уникальных заявок: {total:,}")
    print(f"   ✅ Положительных: {positive:,} ({positive/total*100:.2f}%)")
    print(f"   ❌ Отрицательных: {negative:,} ({negative/total*100:.2f}%)")
    print(f"   🚫 Закрытых: {closed:,} ({closed/total*100:.2f}%)")
    print(f"   📞 Дозвонились: {dozonil:,} ({dozonil/total*100:.2f}%)")
    if dozonil > 0:
        print(f"   🎯 Конверсия: {efficiency:.2f}% (от дозвонившихся)")
except Exception as e:
    print(f"\n❌ ОШИБКА при сохранении: {e}")
    sys.exit(1)
